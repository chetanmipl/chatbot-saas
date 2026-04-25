# backend/app/services/document_service.py
import os, re, uuid
import numpy as np
import aiofiles
from pathlib import Path
from fastapi import UploadFile, HTTPException
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from langchain_text_splitters import RecursiveCharacterTextSplitter  # Level 2
from langchain_text_splitters import MarkdownHeaderTextSplitter       # Level 3

from app.models.document import Document, DocumentStatus
from app.models.document_chunk import DocumentChunk
from app.models.chatbot import Chatbot
from app.models.tenant import Tenant
from app.ai.embeddings import embed_texts, embed_text
from app.core.config import settings

ALLOWED_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
}

UPLOAD_DIR = Path(settings.UPLOAD_DIR)
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


# ─────────────────────────────────────────────────────────────
# TEXT EXTRACTION
# ─────────────────────────────────────────────────────────────

def extract_text(file_path: Path, file_type: str) -> str:
    if file_type == "pdf":
        from pypdf import PdfReader
        reader = PdfReader(str(file_path))
        pages  = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            pages.append(f"[Page {i+1}]\n{text}")
        return "\n\n".join(pages)

    elif file_type == "docx":
        from docx import Document as DocxDocument
        doc   = DocxDocument(str(file_path))
        parts = []
        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "")
                parts.append(f"\n{'#' * int(level)} {para.text}\n")
            elif para.text.strip():
                parts.append(para.text)
        return "\n".join(parts)

    elif file_type == "txt":
        return file_path.read_text(encoding="utf-8")

    raise HTTPException(status_code=400, detail=f"Unsupported: {file_type}")


def clean_text(text: str) -> str:
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    text = text.replace('\xa0', ' ').replace('\u2013', '-').replace('\u2014', '-')
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r' {2,}', ' ', text)
    text = re.sub(r'^\s*[\d\W]{1,3}\s*$', '', text, flags=re.MULTILINE)
    return text.strip()


# ─────────────────────────────────────────────────────────────
# LEVEL 2 — Recursive Character Splitting (LangChain built-in)
# Fallback for dense/narrative text
# ─────────────────────────────────────────────────────────────

def level2_recursive_chunk(text: str) -> list[dict]:
    """
    LangChain's RecursiveCharacterTextSplitter.
    Tries: paragraph → sentence → word boundaries in order.
    Never cuts mid-sentence.
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=600,
        chunk_overlap=120,
        separators=["\n\n", "\n", ". ", "? ", "! ", " ", ""],
        length_function=len,
    )
    docs = splitter.create_documents([text])
    return [
        {"text": d.page_content, "section": "content", "type": "recursive"}
        for d in docs
        if len(d.page_content.split()) >= 10
    ]


# ─────────────────────────────────────────────────────────────
# LEVEL 3 — Document Structure Splitting
# Uses regex patterns + CONTENT_PATTERNS exclusion
# ─────────────────────────────────────────────────────────────

SECTION_PATTERNS = [
    r"^#{1,4}\s+.+",
    r"^Article\s+\d+",
    r"^Section\s+[\d\.]+\s",
    r"^Chapter\s+\d+",
    r"^Clause\s+\d+",
    r"^Part\s+[IVXivx\d]+",
    r"^Schedule\s+\d+",
    r"^Appendix\s+[A-Z\d]+",
    r"^\d+\.\d+[\.\d]*\s+[A-Z].{5,}",  # 3.3 Description — NOT "1. Do this"
    r"^Q:\s+.+",
    r"^[A-Z][A-Z\s]{10,}$",
    r"^\[Page\s+\d+\]",
]

# These LOOK like headings but are content — never split on them
CONTENT_PATTERNS = [
    r"^Phase\s+\d+",
    r"^Step\s+\d+\s*:",
    r"^Stage\s+\d+",
    r"^Round\s+\d+",
    r"^Day\s+\d+",
    r"^Week\s+\d+",
    r"^\d+\.\s+[A-Z]",   # "1. Do this" — list items
]

SECTION_RE = re.compile("|".join(SECTION_PATTERNS), re.MULTILINE)
CONTENT_RE = re.compile("|".join(CONTENT_PATTERNS), re.MULTILINE | re.IGNORECASE)


def level3_structure_chunk(text: str, max_chunk_size: int = 800) -> list[dict]:
    """
    Splits on structural headings detected by regex.
    Skips lines that look like headings but are actually list items or sub-phases.
    """
    lines           = text.split("\n")
    chunks          = []
    current_title   = "Introduction"
    current_content = []

    def flush(title, content):
        full_text = "\n".join(content).strip()
        if not full_text or len(full_text.split()) < 10:
            return []
        if len(full_text) <= max_chunk_size:
            return [{"text": f"{title}\n\n{full_text}",
                     "section": title, "type": "structure"}]
        # Section too large — fall back to Level 2 inside it
        return level2_recursive_chunk(full_text)

    for line in lines:
        stripped        = line.strip()
        is_content_line = bool(CONTENT_RE.match(stripped))
        is_heading      = bool(SECTION_RE.match(stripped)) and not is_content_line

        current_text = "\n".join(current_content).strip()
        has_content  = len(current_text.split()) >= 15

        if is_heading and has_content:
            chunks.extend(flush(current_title, current_content))
            current_title   = stripped
            current_content = []
        else:
            current_content.append(line)

    chunks.extend(flush(current_title, current_content))
    return chunks


# ─────────────────────────────────────────────────────────────
# LEVEL 4 — Semantic Splitting
# Embeds every sentence, splits where meaning shifts sharply
# ─────────────────────────────────────────────────────────────

async def level4_semantic_chunk(
    text: str,
    breakpoint_threshold: float = 0.35,
) -> list[dict]:
    """
    Greg Kamradt's semantic chunking:
    1. Split into sentences
    2. Embed each sentence
    3. Calculate cosine similarity between adjacent sentences
    4. Where similarity drops below threshold = chunk boundary
    """
    # Split into sentences
    sentences = re.split(r"(?<=[.!?])\s+(?=[A-Z])", text.strip())
    sentences = [s.strip() for s in sentences if len(s.strip().split()) > 5]

    if len(sentences) < 4:
        return level2_recursive_chunk(text)

    print(f"  🧠 Semantic chunking: embedding {len(sentences)} sentences...")

    # Embed all sentences at once (one batch call — efficient)
    embeddings = await embed_texts(sentences, is_query=False)
    embeddings = [np.array(e) for e in embeddings]

    # Calculate similarity between adjacent sentences
    similarities = []
    for i in range(len(embeddings) - 1):
        a, b = embeddings[i], embeddings[i + 1]
        sim  = float(np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b) + 1e-8))
        similarities.append(sim)

    # Find breakpoints — where similarity drops sharply
    if similarities:
        mean_sim = np.mean(similarities)
        std_sim  = np.std(similarities)
        # Breakpoint = similarity drops more than 1 std dev below mean
        threshold = mean_sim - (std_sim * breakpoint_threshold)
    else:
        threshold = 0.3

    breakpoints = [
        i for i, sim in enumerate(similarities)
        if sim < threshold
    ]

    print(f"  📍 Found {len(breakpoints)} semantic breakpoints "
          f"(threshold={threshold:.3f})")

    # Build chunks from breakpoints
    chunks  = []
    start   = 0

    for bp in breakpoints:
        chunk_sentences = sentences[start:bp + 1]
        chunk_text      = " ".join(chunk_sentences)
        if len(chunk_text.split()) >= 10:
            chunks.append({
                "text":    chunk_text,
                "section": chunk_sentences[0][:60],
                "type":    "semantic"
            })
        start = bp + 1

    # Last chunk
    if start < len(sentences):
        chunk_text = " ".join(sentences[start:])
        if len(chunk_text.split()) >= 10:
            chunks.append({
                "text":    chunk_text,
                "section": sentences[start][:60],
                "type":    "semantic"
            })

    # If semantic chunking produced too few or too many — fallback
    if len(chunks) < 2 or len(chunks) > 100:
        print(f"  ↩️  Semantic chunking result poor ({len(chunks)} chunks) → using recursive")
        return level2_recursive_chunk(text)

    return chunks


# ─────────────────────────────────────────────────────────────
# SMART ROUTER — picks level per document automatically
# ─────────────────────────────────────────────────────────────

def detect_document_type(text: str) -> str:
    heading_count = len(SECTION_RE.findall(text))
    has_articles  = bool(re.search(r"article\s+\d+", text, re.I))
    has_sections  = bool(re.search(r"section\s+[\d\.]+", text, re.I))
    has_clauses   = bool(re.search(r"clause\s+\d+", text, re.I))
    has_faq       = bool(re.search(r"^q\s*:", text, re.I | re.M))
    has_numbered  = bool(re.search(r"^\d+\.\d+\s+[A-Z]", text, re.M))
    has_markdown  = bool(re.search(r"^#{1,4}\s+", text, re.M))

    paras         = [p for p in text.split("\n\n") if p.strip()]
    avg_para_len  = sum(len(p) for p in paras) / max(len(paras), 1)

    if has_articles or has_clauses or has_faq or heading_count > 3 \
            or has_numbered or has_markdown:
        return "structured"
    elif avg_para_len > 400:
        return "dense"
    else:
        return "mixed"


def _chunk_quality_ok(chunks: list[dict]) -> bool:
    """Returns True if Level 3 produced good chunks."""
    if not chunks:
        return False
    avg_words = sum(len(c["text"].split()) for c in chunks) / len(chunks)
    return (
        3 <= len(chunks) <= 150   # reasonable number
        and avg_words >= 20        # not heading-only chunks
        and avg_words <= 400       # not massive walls of text
    )


async def smart_chunk(text: str) -> list[dict]:
    """
    Routes to the right level:
    Structured doc → Level 3 (structure) → Level 2 fallback
    Dense doc      → Level 4 (semantic)  → Level 2 fallback
    Mixed          → Level 3 → quality check → Level 4 if poor
    """
    doc_type = detect_document_type(text)
    print(f"  📄 Document type: {doc_type}")

    if doc_type == "structured":
        # Level 3 primary
        chunks = level3_structure_chunk(text)

        if not _chunk_quality_ok(chunks):
            print(f"  ⚠️  Level 3 quality poor → trying Level 4 semantic")
            chunks = await level4_semantic_chunk(text)

        if not _chunk_quality_ok(chunks):
            print(f"  ↩️  Level 4 poor → Level 2 recursive fallback")
            chunks = level2_recursive_chunk(text)

    elif doc_type == "dense":
        # Level 4 primary for dense text — finds meaning shifts
        chunks = await level4_semantic_chunk(text)

        if not _chunk_quality_ok(chunks):
            print(f"  ↩️  Semantic poor → Level 2 recursive fallback")
            chunks = level2_recursive_chunk(text)

    else:  # mixed
        # Try Level 3 first, semantic if poor
        chunks = level3_structure_chunk(text)

        if not _chunk_quality_ok(chunks):
            chunks = await level4_semantic_chunk(text)

        if not _chunk_quality_ok(chunks):
            chunks = level2_recursive_chunk(text)

    chunks = [c for c in chunks if len(c["text"].split()) >= 10]

    avg_w = sum(len(c["text"].split()) for c in chunks) / max(len(chunks), 1)
    print(f"  ✂️  {len(chunks)} chunks | avg {avg_w:.0f} words | "
          f"types: {set(c['type'] for c in chunks)}")

    return chunks


# ─────────────────────────────────────────────────────────────
# UPLOAD — unchanged interface, better chunking underneath
# ─────────────────────────────────────────────────────────────

def normalize_embedding(embedding: list[float]) -> list[float]:
    vec  = np.array(embedding, dtype=np.float32)
    norm = np.linalg.norm(vec)
    return (vec / norm).tolist() if norm > 0 else embedding


async def upload_document(
    file: UploadFile,
    chatbot_id: str,
    tenant: Tenant,
    db: AsyncSession,
) -> Document:

    if file.content_type not in ALLOWED_TYPES:
        raise HTTPException(status_code=400, detail="Only PDF, DOCX, TXT supported")

    file_type = ALLOWED_TYPES[file.content_type]

    result = await db.execute(
        select(Chatbot).where(
            Chatbot.id == chatbot_id,
            Chatbot.tenant_id == tenant.id
        )
    )
    chatbot = result.scalar_one_or_none()
    if not chatbot:
        raise HTTPException(status_code=404, detail="Chatbot not found")

    file_id   = str(uuid.uuid4())
    save_path = UPLOAD_DIR / f"{file_id}_{file.filename}"

    async with aiofiles.open(save_path, "wb") as f:
        content = await file.read()
        await f.write(content)

    document = Document(
        filename=file.filename,
        file_type=file_type,
        file_size=len(content),
        file_path=str(save_path),
        status=DocumentStatus.PROCESSING,
        chatbot_id=chatbot_id,
        tenant_id=tenant.id,
    )
    db.add(document)
    await db.flush()

    try:
        raw_text = extract_text(save_path, file_type)
        raw_text = clean_text(raw_text)

        if not raw_text.strip():
            raise ValueError("No text extracted")

        # Smart routing across levels
        chunks = await smart_chunk(raw_text)

        # Embed document chunks (is_query=False — no BGE prefix)
        texts_to_embed = [c["text"] for c in chunks]
        embeddings     = await embed_texts(texts_to_embed, is_query=False)
        embeddings     = [normalize_embedding(e) for e in embeddings]

        for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
            db_chunk = DocumentChunk(
                content=chunk["text"],
                chunk_index=i,
                filename=file.filename,
                embedding=embedding,
                document_id=document.id,
                chatbot_id=chatbot_id,
                tenant_id=tenant.id,
            )
            db.add(db_chunk)

        document.status      = DocumentStatus.READY
        document.chunk_count = len(chunks)
        print(f"  ✅ {file.filename} → {len(chunks)} chunks stored")

    except Exception as e:
        document.status    = DocumentStatus.FAILED
        document.error_msg = str(e)
        print(f"  ❌ Failed: {e}")

    await db.commit()
    await db.refresh(document)
    return document