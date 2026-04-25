# backend/app/services/document_service.py
import os
import re
import uuid
import aiofiles
from pathlib import Path
from fastapi import UploadFile, HTTPException
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from app.models.document import Document, DocumentStatus
from app.models.document_chunk import DocumentChunk
from app.models.chatbot import Chatbot
from app.models.tenant import Tenant
from app.ai.embeddings import embed_texts
from app.core.config import settings
from app.ai.rag_pipeline import llm
from langchain_core.messages import HumanMessage, SystemMessage
import json as _json
import numpy as np

ALLOWED_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
}

UPLOAD_DIR = Path(settings.UPLOAD_DIR)
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


# ─────────────────────────────────────────────────────────────
# EXTRACTION
# ─────────────────────────────────────────────────────────────

def extract_text(file_path: Path, file_type: str) -> str:
    if file_type == "pdf":
        from pypdf import PdfReader
        reader = PdfReader(str(file_path))
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            # Keep page markers — helps with "see page 3" queries
            pages.append(f"[Page {i+1}]\n{text}")
        return "\n\n".join(pages)

    elif file_type == "docx":
        from docx import Document as DocxDocument
        doc = DocxDocument(str(file_path))
        parts = []
        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                parts.append(f"\n## {para.text}\n")
            elif para.text.strip():
                parts.append(para.text)
        return "\n".join(parts)

    elif file_type == "txt":
        return file_path.read_text(encoding="utf-8")

    raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_type}")


# ─────────────────────────────────────────────────────────────
# TIER 1 — SEMANTIC SECTION CHUNKING
# Best for: legal docs, articles, manuals, structured reports
# ─────────────────────────────────────────────────────────────

# Patterns that signal a new section in most professional documents
# SECTION_PATTERNS = [
#     # Markdown headings
#     r"^#{1,4}\s+.+",

#     # Legal / Academic
#     r"^Article\s+\d+",
#     r"^Section\s+[\d\.]+",
#     r"^Chapter\s+\d+",
#     r"^Clause\s+\d+",
#     r"^Part\s+[IVXivx\d]+",
#     r"^Schedule\s+\d+",
#     r"^Appendix\s+[A-Z\d]+",

#     # Business documents
#     r"^Q:\s+.+",                          # FAQ format
#     r"^\d+\.\s+[A-Z].{10,}",             # 1. Introduction
#     r"^\d+\.\d+\s+[A-Z].{5,}",           # 1.1 Overview
#     r"^[A-Z][a-z]+(\s[A-Z][a-z]+){0,4}:$", # "Refund Policy:" as heading

#     # ALL CAPS headings (common in contracts, manuals)
#     r"^[A-Z][A-Z\s]{8,}$",

#     # HR / Policy docs
#     r"^Policy\s*:\s*.+",
#     r"^Procedure\s*:\s*.+",

#     # Product manuals
#     r"^Step\s+\d+",
#     r"^Installation",
#     r"^Troubleshooting",

#     # Page markers (added during PDF extraction)
#     r"^\[Page\s+\d+\]",
# ]

# SECTION_RE = re.compile("|".join(SECTION_PATTERNS), re.MULTILINE)
SECTION_PATTERNS = [
    # Markdown
    r"^#{1,4}\s+.+",

    # Legal / Academic
    r"^Article\s+\d+",
    r"^Section\s+[\d\.]+\s",      # "Section 2.3 " — note trailing space
    r"^Chapter\s+\d+",
    r"^Clause\s+\d+",
    r"^Part\s+[IVXivx\d]+",
    r"^Schedule\s+\d+",
    r"^Appendix\s+[A-Z\d]+",

    # Numbered sections X.X format ONLY — not list items
    # "3.3 Description" ✅   "1. Do this" ❌
    r"^\d+\.\d+[\.\d]*\s+[A-Z].{5,}",

    # Business
    r"^Q:\s+.+",
    r"^[A-Z][a-z]+(\s[A-Z][a-z]+){0,4}:$",

    # ALL CAPS (min 3 words to avoid false positives)
    r"^[A-Z][A-Z\s]{10,}$",

    # Page markers from PDF extraction
    r"^\[Page\s+\d+\]",
]

# Lines that look like headings but are actually CONTENT
# Never split on these — they belong to their parent section
CONTENT_PATTERNS = [
    r"^Phase\s+\d+",           # Phase 1: Pre-Event
    r"^Step\s+\d+\s*:",        # Step 1: Setup
    r"^Stage\s+\d+",
    r"^Round\s+\d+",
    r"^Day\s+\d+",
    r"^Week\s+\d+",
    r"^\d+\.\s+[A-Z]",         # "1. Do this" — numbered LIST ITEMS
]

SECTION_RE  = re.compile("|".join(SECTION_PATTERNS),  re.MULTILINE)
CONTENT_RE  = re.compile("|".join(CONTENT_PATTERNS),  re.MULTILINE | re.IGNORECASE)



# def semantic_chunk(text: str, max_chunk_size: int = 600) -> list[dict]:
#     """
#     Split text at natural section boundaries (headings, articles, clauses).
#     Each chunk keeps its section title for context.
#     Returns list of {text, section_title, chunk_type}
#     """
#     lines  = text.split("\n")
#     chunks = []

#     current_title   = "Introduction"
#     current_content = []

#     def flush(title: str, content: list[str]) -> list[dict]:
#         """Save current section, split further if too large."""
#         full_text = "\n".join(content).strip()
#         if not full_text:
#             return []

#         if len(full_text) <= max_chunk_size:
#             return [{"text": f"{title}\n\n{full_text}", "section": title, "type": "section"}]

#         # Section too large — split into sentence windows
#         return sentence_window_chunk(full_text, title=title,
#                                      chunk_size=max_chunk_size, overlap=80)

#     for line in lines:
#         is_heading = SECTION_RE.match(line.strip())

#         if is_heading and current_content:
#             # Save previous section
#             chunks.extend(flush(current_title, current_content))
#             current_title   = line.strip()
#             current_content = []
#         else:
#             current_content.append(line)

#     # Don't forget last section
#     chunks.extend(flush(current_title, current_content))
#     return chunks

def semantic_chunk(text: str, max_chunk_size: int = 800) -> list[dict]:
    lines           = text.split("\n")
    chunks          = []
    current_title   = "Introduction"
    current_content = []

    def flush(title, content):
        full_text = "\n".join(content).strip()
        if not full_text:
            return []
        if len(full_text) <= max_chunk_size:
            return [{"text": f"{title}\n\n{full_text}",
                     "section": title, "type": "section"}]
        return sentence_window_chunk(full_text, title=title,
                                     chunk_size=max_chunk_size, overlap=120)

    for line in lines:
        stripped = line.strip()

        is_content_line = bool(CONTENT_RE.match(stripped))  # check first
        is_heading      = bool(SECTION_RE.match(stripped)) and not is_content_line

        current_text = "\n".join(current_content).strip()
        has_content  = len(current_text.split()) >= 15

        if is_heading and has_content:
            chunks.extend(flush(current_title, current_content))
            current_title   = stripped
            current_content = []
        else:
            # content line, list item, sub-heading — stays with parent
            current_content.append(line)

    chunks.extend(flush(current_title, current_content))
    return chunks

# ─────────────────────────────────────────────────────────────
# TIER 2 — SENTENCE WINDOW CHUNKING
# Best for: dense paragraphs, narrative text, PDFs
# ─────────────────────────────────────────────────────────────

def sentence_window_chunk(
    text: str,
    title: str = "",
    chunk_size: int = 400,
    overlap: int = 100,
) -> list[dict]:
    """
    Splits on sentence boundaries with overlap.
    Overlap carries WHOLE sentences, not arbitrary chars —
    so no sentence is cut in half at a boundary.
    """
    # Split into sentences
    sentences = re.split(r"(?<=[.!?])\s+(?=[A-Z])", text.strip())
    sentences = [s.strip() for s in sentences if len(s.strip()) > 15]

    chunks    = []
    window    = []
    win_len   = 0

    for sent in sentences:
        if win_len + len(sent) > chunk_size and window:
            chunk_text = " ".join(window)
            if title:
                chunk_text = f"{title}\n\n{chunk_text}"
            chunks.append({
                "text":    chunk_text,
                "section": title,
                "type":    "sentence_window"
            })

            # Carry overlap: keep sentences from the end until overlap is filled
            overlap_sents = []
            overlap_len   = 0
            for s in reversed(window):
                if overlap_len + len(s) <= overlap:
                    overlap_sents.insert(0, s)
                    overlap_len += len(s)
                else:
                    break

            window  = overlap_sents + [sent]
            win_len = sum(len(s) for s in window)
        else:
            window.append(sent)
            win_len += len(sent)

    if window:
        chunk_text = " ".join(window)
        if title:
            chunk_text = f"{title}\n\n{chunk_text}"
        chunks.append({
            "text":    chunk_text,
            "section": title,
            "type":    "sentence_window"
        })

    return chunks

#Clean Chunks
def clean_text(text: str) -> str:
    """Remove garbage before chunking — critical for PDF quality."""
    # Remove null bytes and non-printable chars
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    # Normalize unicode spaces and dashes
    text = text.replace('\xa0', ' ').replace('\u2013', '-').replace('\u2014', '-')
    # Collapse 3+ newlines into 2
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Collapse multiple spaces
    text = re.sub(r' {2,}', ' ', text)
    # Remove lines that are just numbers or single chars (PDF artifacts)
    text = re.sub(r'^\s*[\d\W]{1,3}\s*$', '', text, flags=re.MULTILINE)
    return text.strip()


# ─────────────────────────────────────────────────────────────
# TIER 3 — PARENT-CHILD CHUNKING
# Best for: when you want precise retrieval but full context in answers
# Retrieve small child chunks, but send parent chunk to LLM
# ─────────────────────────────────────────────────────────────

def parent_child_chunk(chunks: list[dict]) -> list[dict]:
    """
    For each section chunk (parent), create smaller child chunks for retrieval.
    The child stores the parent text so the LLM gets full context.
    """
    result = []

    for parent in chunks:
        parent_text = parent["text"]

        # If parent is small enough, use as-is
        if len(parent_text) <= 250:
            result.append(parent)
            continue

        # Split parent into small child chunks (for precise embedding)
        sentences  = re.split(r"(?<=[.!?])\s+", parent_text)
        child_size = 150
        child_buf  = []
        child_len  = 0

        for sent in sentences:
            if child_len + len(sent) > child_size and child_buf:
                child_text = " ".join(child_buf)
                result.append({
                    # Small text for embedding (precise retrieval)
                    "text":        child_text,
                    # Full parent text sent to LLM (full context)
                    "parent_text": parent_text,
                    "section":     parent["section"],
                    "type":        "child",
                })
                child_buf = [sent]
                child_len = len(sent)
            else:
                child_buf.append(sent)
                child_len += len(sent)

        if child_buf:
            result.append({
                "text":        " ".join(child_buf),
                "parent_text": parent_text,
                "section":     parent["section"],
                "type":        "child",
            })

    return result


# ─────────────────────────────────────────────────────────────
# SMART ROUTER — picks the right strategy per document
# ─────────────────────────────────────────────────────────────

def detect_document_type(text: str) -> str:
    text_lower     = text.lower()
    heading_count  = len(SECTION_RE.findall(text))
    has_articles   = bool(re.search(r"article\s+\d+", text_lower))
    has_sections   = bool(re.search(r"section\s+[\d\.]+", text_lower))
    has_clauses    = bool(re.search(r"clause\s+\d+", text_lower))
    has_faq        = bool(re.search(r"^q\s*:", text_lower, re.MULTILINE))
    has_numbered   = bool(re.search(r"^\d+\.\s+[A-Z]", text, re.MULTILINE))

    paras          = [p for p in text.split("\n\n") if p.strip()]
    avg_para_len   = sum(len(p) for p in paras) / max(len(paras), 1)

    if has_articles or has_clauses or has_faq or heading_count > 3 or has_numbered:
        return "structured"
    elif avg_para_len > 400:
        return "dense"
    else:
        return "mixed"       # everything else → semantic + sentence window


async def detect_headings_with_llm(text: str) -> list[str]:
    """
    Give LLM a sample of the document, ask it to identify
    what heading patterns exist. Returns regex patterns
    specific to THIS document — not hardcoded.
    """


    # Send first 3000 chars — enough to detect structure
    sample = text[:3000]

    prompt = [
        SystemMessage(content="""Analyze this document sample and identify ALL heading/section patterns.

Return a JSON object with:
{
  "heading_lines": ["exact line 1", "exact line 2", ...],  // actual heading lines you see
  "pattern_description": "brief description of heading style"
}

A heading is any line that:
- Introduces a new section or topic
- Is shorter than surrounding content
- May be numbered (3.3, Phase 1, Article 14)
- May be bold or ALL CAPS in the original

Output ONLY valid JSON, nothing else."""),
        HumanMessage(content=f"Document sample:\n\n{sample}")
    ]

    response = await llm.ainvoke(prompt)

    try:
        raw  = response.content.strip()
        raw  = raw[raw.find("{"):raw.rfind("}") + 1]
        data = _json.loads(raw)
        heading_lines = data.get("heading_lines", [])
        print(f"  🤖 LLM detected {len(heading_lines)} heading patterns")
        return heading_lines
    except Exception as e:
        print(f"  ⚠️ LLM heading detection failed: {e} — using fallback")
        return []
    
def llm_guided_chunk(
    text: str,
    heading_lines: list[str],
    max_chunk_size: int = 800
) -> list[dict]:
    """
    Split document using headings the LLM identified.
    Works for ANY document type — no patterns needed.
    """
    if not heading_lines:
        return sentence_window_chunk(text, chunk_size=max_chunk_size)

    lines   = text.split("\n")
    chunks  = []
    current_title   = "Introduction"
    current_content = []

    # Build a lookup set of heading lines for fast matching
    heading_set = set(h.strip().lower() for h in heading_lines)

    def flush(title, content):
        full_text = "\n".join(content).strip()
        if not full_text or len(full_text.split()) < 10:
            return []
        if len(full_text) <= max_chunk_size:
            return [{"text": f"{title}\n\n{full_text}",
                     "section": title, "type": "llm_guided"}]
        return sentence_window_chunk(full_text, title=title,
                                     chunk_size=max_chunk_size, overlap=120)

    for line in lines:
        stripped  = line.strip()
        is_heading = stripped.lower() in heading_set

        # Also check partial match for long headings
        if not is_heading and stripped:
            is_heading = any(
                stripped.lower().startswith(h[:30].lower())
                for h in heading_set if len(h) > 10
            )

        current_text = "\n".join(current_content).strip()
        has_content  = len(current_text.split()) >= 15

        if is_heading and has_content:
            chunks.extend(flush(current_title, current_content))
            current_title   = stripped
            current_content = []
        else:
            current_content.append(line)

    chunks.extend(flush(current_title, current_content))
    return chunks

# def smart_chunk(text: str) -> list[dict]:
#     """
#     Automatically picks and applies the best chunking strategy.
#     This is what runs on every document upload.
#     """
#     doc_type = detect_document_type(text)
#     print(f"  📄 Document type detected: {doc_type}")

#     if doc_type == "structured":
#         # Semantic sections → then parent-child for precise retrieval
#         sections = semantic_chunk(text, max_chunk_size=600)
#         chunks   = parent_child_chunk(sections)

#     elif doc_type == "dense":
#         # Pure sentence windows with generous overlap
#         chunks = sentence_window_chunk(text, chunk_size=350, overlap=120)

#     else:  # mixed
#         # Semantic first, sentence window fallback inside large sections
#         chunks = semantic_chunk(text, max_chunk_size=500)

#     # Final cleanup — remove empty or tiny chunks
#     chunks = [
#         c for c in chunks
#         if len(c["text"].split()) >= 10
#     ]

#     print(f"  ✂️  Created {len(chunks)} chunks using {doc_type} strategy")
#     return chunks
async def smart_chunk(text: str) -> list[dict]:
    doc_type = detect_document_type(text)

    if doc_type == "dense":
        # No headings expected — sentence window directly
        return sentence_window_chunk(text, chunk_size=400, overlap=120)

    # Try regex first — fast, private, free
    chunks = semantic_chunk(text, max_chunk_size=800)

    # ── Quality check ─────────────────────────────────────────
    avg_chunk_words = sum(len(c["text"].split()) for c in chunks) / max(len(chunks), 1)
    too_many_chunks = len(chunks) > 80       # over-split
    too_few_chunks  = len(chunks) < 3        # under-split
    chunks_too_tiny = avg_chunk_words < 20   # heading-only chunks

    needs_llm = too_many_chunks or too_few_chunks or chunks_too_tiny

    if needs_llm:
        # Only now use LLM — something went wrong with regex
        print(f"  ⚠️  Regex chunking poor quality "
              f"({len(chunks)} chunks, avg {avg_chunk_words:.0f} words) "
              f"→ falling back to LLM detection")
        heading_lines = await detect_headings_with_llm(text)
        if heading_lines:
            chunks = llm_guided_chunk(text, heading_lines, max_chunk_size=800)

    chunks = [c for c in chunks if len(c["text"].split()) >= 10]
    print(f"  ✂️  {len(chunks)} chunks, avg {avg_chunk_words:.0f} words/chunk")
    return chunks

def normalize_embedding(embedding: list[float]) -> list[float]:
    """Normalize to unit vector so cosine similarity works correctly."""
    vec  = np.array(embedding, dtype=np.float32)
    norm = np.linalg.norm(vec)
    if norm == 0:
        return embedding
    return (vec / norm).tolist()


# ─────────────────────────────────────────────────────────────
# MAIN UPLOAD FUNCTION
# ─────────────────────────────────────────────────────────────

async def upload_document(
    file: UploadFile,
    chatbot_id: str,
    tenant: Tenant,
    db: AsyncSession,
) -> Document:

    if file.content_type not in ALLOWED_TYPES:
        raise HTTPException(status_code=400, detail="Only PDF, DOCX, TXT supported")

    file_type = ALLOWED_TYPES[file.content_type]

    result = await db.execute(
        select(Chatbot).where(Chatbot.id == chatbot_id, Chatbot.tenant_id == tenant.id)
    )
    chatbot = result.scalar_one_or_none()
    if not chatbot:
        raise HTTPException(status_code=404, detail="Chatbot not found")

    file_id   = str(uuid.uuid4())
    save_path = UPLOAD_DIR / f"{file_id}_{file.filename}"

    async with aiofiles.open(save_path, "wb") as f:
        content = await file.read()
        await f.write(content)

    document = Document(
        filename=file.filename,
        file_type=file_type,
        file_size=len(content),
        file_path=str(save_path),
        status=DocumentStatus.PROCESSING,
        chatbot_id=chatbot_id,
        tenant_id=tenant.id,
    )
    db.add(document)
    await db.flush()

    try:
        # Extract text
        raw_text = extract_text(save_path, file_type)
        if not raw_text.strip():
            raise ValueError("No text extracted from file")
        
        #clean chunk
        raw_text = clean_text(raw_text) 

        # Smart chunk
        chunks = await smart_chunk(raw_text)

        # Embed — use parent_text if available (child chunking), else text
        texts_to_embed = [
            c.get("parent_text", c["text"]) for c in chunks
        ]
        embeddings = await embed_texts(texts_to_embed)
        embeddings = [normalize_embedding(e) for e in embeddings]

        # Store chunks
        for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
            db_chunk = DocumentChunk(
                content=chunk.get("parent_text", chunk["text"]),  # full context for LLM
                chunk_index=i,
                filename=file.filename,
                embedding=embedding,
                document_id=document.id,
                chatbot_id=chatbot_id,
                tenant_id=tenant.id,
            )
            db.add(db_chunk)

        document.status      = DocumentStatus.READY
        document.chunk_count = len(chunks)
        print(f"  ✅ Document processed: {len(chunks)} chunks stored")

    except Exception as e:
        document.status    = DocumentStatus.FAILED
        document.error_msg = str(e)
        print(f"  ❌ Document processing failed: {e}")

    await db.commit()
    await db.refresh(document)
    return document